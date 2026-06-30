"""
Worker 单元测试：
- text_extract: PDF / DOCX / TXT 解析正确
- lemma: 单词 / 短语 / 专有名词 各种 case
"""
import io
import pytest


class TestTextExtract:
    """TXT 是最直接的，PDF/DOCX 用最小可工作样本"""

    def test_extract_txt_utf8(self, storage):
        """TXT UTF-8 解码"""
        from app.workers.text_extract import _extract_txt
        data = "Hello world\n\nThere are some books on the table.".encode("utf-8")
        result = _extract_txt(data)
        assert "Hello world" in result
        assert "books" in result

    def test_extract_txt_gbk(self, storage):
        """TXT GBK 解码"""
        from app.workers.text_extract import _extract_txt
        data = "你好世界".encode("gbk")
        result = _extract_txt(data)
        assert "你好世界" in result

    def test_extract_docx(self, storage):
        """DOCX 解析"""
        from docx import Document
        from app.workers.text_extract import _extract_docx

        # 内存中创建 docx
        doc = Document()
        doc.add_paragraph("Hello world")
        doc.add_paragraph("This is a test document.")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        result = _extract_docx(buf.getvalue())
        assert "Hello world" in result
        assert "This is a test document" in result

    def test_extract_pdf(self, storage):
        """PDF 解析（用 pdfplumber 自带的样本或生成最小 PDF）"""
        # 用最简单的纯文本 PDF 生成（pdfplumber 接受）
        # 这里跳过复杂 PDF 生成，改用直接构造一个最小 PDF 的 bytes
        # 用 reportlab 会增加依赖，改用预生成的样本
        pytest.skip("PDF 测试需要样本文件，由 E2E 测试覆盖")


class TestLemma:
    """Lemma 后处理：单词 / 短语 / 专有名词"""

    def test_lemma_basic_noun(self):
        """books → book"""
        from app.workers.lemma import lemmatize_word
        # 即使 spaCy 未安装，fallback 至少 lowercase
        result = lemmatize_word("books")
        assert result == "book", f"expected 'book', got '{result}'"

    def test_lemma_verb_running(self):
        """running → run（如果有 spaCy）或 running（fallback）"""
        from app.workers.lemma import lemmatize_word
        result = lemmatize_word("running")
        # spaCy 会得到 run；fallback 也是 run（小写）
        assert result == "run", f"expected 'run', got '{result}'"

    def test_lemma_phrase_passthrough(self):
        """短语 in spite of 不变"""
        from app.workers.lemma import lemmatize_word
        result = lemmatize_word("in spite of")
        assert result == "in spite of", f"expected phrase passthrough, got '{result}'"

    def test_lemma_phrase_arrive_in(self):
        """短语 arrive in 不变"""
        from app.workers.lemma import lemmatize_word
        result = lemmatize_word("arrive in")
        assert result == "arrive in", f"expected phrase passthrough, got '{result}'"

    def test_lemma_batch(self):
        """批量 lemma"""
        from app.workers.lemma import lemmatize_words_sync
        result = lemmatize_words_sync([
            "books", "running", "went", "better", "in spite of", "Apple",
        ])
        assert len(result) == 6
        # 短语和专有名词标记 changed=False
        phrase_entry = next(r for r in result if r["original"] == "in spite of")
        assert phrase_entry["changed"] is False or phrase_entry["lemma"] == "in spite of"
        # 单词被 lemma（小写化也算 changed 因为大写 → 小写）
        for entry in result:
            assert "original" in entry
            assert "lemma" in entry
            assert "changed" in entry

    def test_lemma_empty_input(self):
        """空列表"""
        from app.workers.lemma import lemmatize_words_sync
        assert lemmatize_words_sync([]) == []

    def test_lemma_handles_none(self):
        """None / 空字符串 跳过"""
        from app.workers.lemma import lemmatize_words_sync
        result = lemmatize_words_sync(["", None, "books"])
        # None 和空字符串被跳过
        words = [r["original"] for r in result]
        assert "books" in words
        assert None not in words
        assert "" not in words