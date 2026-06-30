"""
文本提取 worker（task #16）：
- .pdf → pdfplumber
- .docx → python-docx
- .txt → 直接读

输入：job_id（从 storage_key 取文件）
输出：Job.result['text'] = 纯文本
后续：调用 lemma worker 做后处理（如已有 LLM 输出）
"""
import io
import logging
from datetime import datetime

from ..celery_app import celery_app
from ..storage import get_storage
from ..database import SessionLocal
from ..models import Job, JobStatus, JobStage
from .pipeline import update_job_status

logger = logging.getLogger(__name__)


def _extract_pdf(data: bytes) -> str:
    """pdfplumber 提取文本"""
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
    return "\n\n".join(text_parts)


def _extract_docx(data: bytes) -> str:
    """python-docx 提取文本（段落 + 表格）"""
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # 表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_txt(data: bytes) -> str:
    """纯文本直接解码"""
    for encoding in ("utf-8", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_image(data: bytes) -> str:
    """OCR 提取图片中的文字（task #15 / task #59）

    使用 EasyOCR（本地引擎），输出与 _extract_pdf/docx/txt 相同的纯文本格式。
    首次调用会触发模型下载（~140MB，缓存到 ~/.EasyOCR/model/）。
    """
    from .ocr_extract import extract_image_sync
    return extract_image_sync(data)


def extract_text_sync(job_id: str) -> str:
    """同步执行文本提取（worker 内部 / 测试可直接调用）"""
    with SessionLocal() as db:
        job = db.query(Job).filter(Job.id == job_id).first()
        if not job:
            raise ValueError(f"Job {job_id} 不存在")
        storage_key = job.storage_key
        source_type = job.source_type

    storage = get_storage()
    data = storage.download(storage_key)

    if source_type == "pdf":
        text = _extract_pdf(data)
    elif source_type == "docx":
        text = _extract_docx(data)
    elif source_type == "txt":
        text = _extract_txt(data)
    elif source_type == "image":
        text = _extract_image(data)
    else:
        raise ValueError(f"不支持的 source_type: {source_type}")

    update_job_status(
        job_id,
        current_stage=JobStage.TEXT_EXTRACT,
        progress=30,
        result={"text": text, "text_length": len(text)},
    )
    return text


@celery_app.task(name="app.workers.text_extract.extract_text_task", bind=True, max_retries=2)
def extract_text_task(self, job_id: str):
    """
    Celery task：提取文本。
    成功后自动触发 lemma worker（后续接 LLM 后会先 LLM 后 lemma）。
    """
    logger.info(f"extract_text_task 启动: job={job_id}")
    update_job_status(job_id, status=JobStatus.PROCESSING, current_stage=JobStage.TEXT_EXTRACT, progress=10)

    try:
        text = extract_text_sync(job_id)
        logger.info(f"提取完成: job={job_id} length={len(text)}")
    except Exception as e:
        logger.exception(f"提取失败: job={job_id}")
        update_job_status(job_id, status=JobStatus.FAILED, error_message=f"文本提取失败: {e}")
        # 重试一次
        try:
            self.retry(exc=e, countdown=10)
        except self.MaxRetriesExceededError:
            pass
        raise

    # 后续：调用 LLM（暂缓，task #14）→ lemma（task #18）
    # 当前实现：直接把文本存到 result，pipeline 暂时停在这里
    update_job_status(
        job_id,
        status=JobStatus.COMPLETED,
        current_stage=JobStage.DONE,
        progress=100,
        result={"text_length": len(text)},
    )
    return {"job_id": job_id, "text_length": len(text)}